from __future__ import annotations

import json
import re
import unicodedata
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"


def load_json(name: str) -> list[dict[str, Any]]:
    with (DATA_DIR / name).open("r", encoding="utf-8") as handle:
        return json.load(handle)


def normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFD", value.lower())
    value = "".join(char for char in value if unicodedata.category(char) != "Mn")
    value = re.sub(r"[^a-z0-9\s/.-]", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def tokenize(value: str) -> set[str]:
    stopwords = {
        "toi",
        "cong",
        "ty",
        "doanh",
        "nghiep",
        "co",
        "la",
        "duoc",
        "khong",
        "can",
        "gi",
        "nao",
        "ve",
        "cho",
        "the",
        "hay",
        "neu",
        "mot",
        "cac",
    }
    return {token for token in normalize_text(value).split() if len(token) > 1 and token not in stopwords}


def classify_sme(profile: dict[str, Any]) -> dict[str, Any]:
    industry = profile.get("industry", "")
    employees = int(profile.get("employees", 0) or 0)
    revenue = float(profile.get("revenue_bil", 0) or 0)
    capital = float(profile.get("capital_bil", 0) or 0)
    is_trade_service = industry in {"Phần mềm / AI", "Dịch vụ đổi mới sáng tạo"}

    if employees <= 10 and revenue <= 10 and capital <= 3:
        size = "Siêu nhỏ"
    elif is_trade_service:
        if employees <= 50 and (revenue <= 100 or capital <= 50):
            size = "Nhỏ"
        elif employees <= 100 and (revenue <= 300 or capital <= 100):
            size = "Vừa"
        else:
            size = "Không thuộc DNNVV"
    else:
        if employees <= 100 and (revenue <= 50 or capital <= 20):
            size = "Nhỏ"
        elif employees <= 200 and (revenue <= 200 or capital <= 100):
            size = "Vừa"
        else:
            size = "Không thuộc DNNVV"

    return {
        "size": size,
        "is_sme": size != "Không thuộc DNNVV",
        "basis": "Phân loại demo theo Điều 5 Nghị định 80/2021/NĐ-CP từ lao động, doanh thu và vốn.",
    }


def match_policies(profile: dict[str, Any]) -> list[dict[str, Any]]:
    policies = load_json("policies.json")
    sme = classify_sme(profile)
    results: list[dict[str, Any]] = []

    for policy in policies:
        eligibility = policy["eligibility"]
        score = 30
        reasons: list[str] = []
        gaps: list[str] = []

        if eligibility.get("requires_sme"):
            if sme["is_sme"]:
                score += 22
                reasons.append(f"Đạt tiêu chí {sme['size']} theo Nghị định 80.")
            else:
                score -= 35
                gaps.append("Chưa đạt tiêu chí DNNVV theo dữ liệu hiện tại.")

        if eligibility.get("requires_startup_innovation"):
            if profile.get("startup_innovation"):
                score += 22
                reasons.append("Hồ sơ đánh dấu là startup đổi mới sáng tạo.")
            else:
                score -= 25
                gaps.append("Chưa có yếu tố khởi nghiệp đổi mới sáng tạo.")

        industries = set(eligibility.get("industries", []))
        if profile.get("industry") in industries:
            score += 18
            reasons.append(f"Lĩnh vực {profile.get('industry')} nằm trong nhóm phù hợp.")
        else:
            score -= 12
            gaps.append("Lĩnh vực chưa nằm trong nhóm ưu tiên của chương trình.")

        provinces = set(eligibility.get("provinces", []))
        if "Tất cả" in provinces or profile.get("province") in provinces:
            score += 8
            if "Tất cả" in provinces:
                reasons.append("Chương trình áp dụng toàn quốc.")
            else:
                reasons.append(f"Chương trình áp dụng tại {profile.get('province')}.")
        else:
            score -= 20
            gaps.append(f"Chương trình chỉ áp dụng tại {', '.join(sorted(provinces))}.")

        if policy["id"] == "p_smedf" and float(profile.get("revenue_bil", 0) or 0) > 0:
            score += 8
            reasons.append("Có doanh thu để bắt đầu chuẩn bị phương án vay và hồ sơ tài chính.")

        if policy["id"] == "p_dean844" and profile.get("startup_innovation"):
            score += 8
            reasons.append("Phù hợp để demo luồng checklist và đơn Đề án 844.")

        score = max(0, min(100, score))
        result = dict(policy)
        result.update(
            {
                "score": score,
                "match_level": "Rất phù hợp" if score >= 80 else "Cần rà soát" if score >= 55 else "Chưa ưu tiên",
                "reasons": reasons[:4],
                "gaps": gaps[:3],
            }
        )
        results.append(result)

    return sorted(results, key=lambda item: item["score"], reverse=True)


@dataclass
class Answer:
    text: str
    citations: list[dict[str, str]]
    confidence: str


def _top_corpus(question: str, limit: int = 3) -> list[dict[str, Any]]:
    question_tokens = tokenize(question)
    ranked = []
    for chunk in load_json("corpus.json"):
        haystack = " ".join([chunk["title"], chunk["clause"], " ".join(chunk["tags"]), chunk["text"]])
        chunk_tokens = tokenize(haystack)
        overlap = len(question_tokens & chunk_tokens)
        exact_bonus = sum(1 for tag in chunk["tags"] if normalize_text(tag).replace("_", " ") in normalize_text(question))
        ranked.append((overlap + exact_bonus * 2, chunk))
    ranked.sort(key=lambda pair: pair[0], reverse=True)
    return [chunk for score, chunk in ranked[:limit] if score > 0]


def answer_question(question: str, profile: dict[str, Any] | None = None) -> Answer:
    q = normalize_text(question)

    if "mien toan bo thue" in q or "ngoai corpus" in q:
        return Answer(
            text="Không đủ thông tin trong corpus demo để kết luận doanh nghiệp được miễn toàn bộ thuế. Corpus hiện chỉ có căn cứ chung về hình thức ưu đãi đầu tư và danh mục ngành nghề; quyết định thuế cần thêm văn bản thuế chuyên ngành, mã ngành, dự án đầu tư và thời điểm áp dụng.",
            citations=[],
            confidence="Ngoài corpus",
        )

    chunks = _top_corpus(question)

    if not chunks:
        return Answer(
            text="Không đủ thông tin trong corpus demo để trả lời chắc chắn. Nên bổ sung văn bản gốc hoặc hỏi lại trong phạm vi DNNVV, Đề án 844, SMEDF, ưu đãi đầu tư hoặc chương trình Hà Nội.",
            citations=[],
            confidence="Ngoài corpus",
        )

    citations = [
        {
            "document": chunk["title"],
            "clause": chunk["clause"],
            "status": chunk["status"],
            "source": chunk["source"],
        }
        for chunk in chunks
    ]

    if "startup" in q and "ho tro" in q:
        text = (
            "Startup phần mềm tại Hà Nội nên ưu tiên ba hướng trong MVP: hỗ trợ DNNVV khởi nghiệp sáng tạo theo Nghị định 80, "
            "Đề án 844 cho hoạt động hệ sinh thái/hoàn thiện năng lực, và chương trình địa phương Hà Nội để kết nối cố vấn, đào tạo, sự kiện. "
            "Nếu có dự án đầu tư hoặc hoạt động sản xuất phần mềm rõ ràng, doanh nghiệp có thể rà soát thêm nhóm ưu đãi đầu tư cho CNTT/phần mềm."
        )
    elif "tu van" in q or "so huu tri tue" in q or "thu nghiem" in q:
        text = (
            "DNNVV khởi nghiệp sáng tạo có thể xem xét hỗ trợ tư vấn, sở hữu trí tuệ, tiêu chuẩn đo lường chất lượng, thử nghiệm và hoàn thiện sản phẩm mới. "
            "Hồ sơ nên mô tả rõ sản phẩm, nhu cầu hỗ trợ và căn cứ chứng minh tính đổi mới."
        )
    elif "nghi dinh 80" in q or "nd 80" in q:
        text = (
            "Trong corpus demo, Nghị định 80/2021/NĐ-CP đang được dùng làm căn cứ trung tâm cho phân loại DNNVV và hỗ trợ DNNVV khởi nghiệp sáng tạo. "
            "Khi hiển thị cho người dùng, MVP gắn badge hiệu lực và link nguồn cho từng điều khoản."
        )
    elif "dnnvv" in q or "nho va vua" in q or "nho/vua" in q:
        if profile:
            sme = classify_sme(profile)
            verdict = "thuộc DNNVV" if sme["is_sme"] else "chưa thuộc DNNVV"
            text = (
                f"Theo dữ liệu hồ sơ hiện tại, {profile.get('name', 'doanh nghiệp')} {verdict}, nhóm {sme['size']}. "
                f"Căn cứ demo dùng lao động {profile.get('employees')} người, doanh thu {profile.get('revenue_bil')} tỷ và vốn {profile.get('capital_bil')} tỷ. "
                "Khi nộp thật cần đối chiếu báo cáo tài chính năm trước liền kề và lĩnh vực hoạt động chính."
            )
        else:
            text = (
                "DNNVV được xác định theo lĩnh vực, số lao động bình quân năm, tổng nguồn vốn hoặc doanh thu năm trước liền kề. "
                "Bạn cần nhập lao động, doanh thu, vốn và lĩnh vực để hệ thống kết luận cụ thể."
            )
    elif "844" in q or "de an" in q:
        text = (
            "Đề án 844 phù hợp với startup đổi mới sáng tạo hoặc tổ chức hỗ trợ hệ sinh thái. "
            "Hồ sơ demo nên chuẩn bị thuyết minh nhiệm vụ, mục tiêu, nội dung, sản phẩm, dự toán kinh phí và năng lực thực hiện."
        )
    elif "smedf" in q or "vay" in q or "von" in q:
        text = (
            "SMEDF là hướng phù hợp khi doanh nghiệp là DNNVV và có phương án sản xuất kinh doanh khả thi, minh bạch tài chính. "
            "Startup hoặc doanh nghiệp tham gia chuỗi giá trị nên chuẩn bị báo cáo tài chính, phương án kinh doanh và tài liệu chứng minh tiêu chí DNNVV."
        )
    elif "phan mem" in q or "ai" in q or "cong nghe thong tin" in q or "uu dai dau tu" in q:
        text = (
            "Sản xuất phần mềm, sản phẩm công nghệ thông tin, nội dung số, R&D và một số hoạt động công nghệ cao có thể thuộc nhóm ngành nghề ưu đãi đầu tư. "
            "Kết luận cuối cùng cần rà soát mã ngành, dự án đầu tư và văn bản thuế liên quan tại thời điểm nộp."
        )
    elif "ha noi" in q:
        text = (
            "Với doanh nghiệp startup tại Hà Nội, có thể theo dõi các chương trình địa phương về đào tạo, cố vấn, kết nối đầu tư và sự kiện hệ sinh thái. "
            "Nguồn này trong MVP đang là seed demo, cần xác minh thông báo chính thức trước khi nộp hồ sơ thật."
        )
    else:
        snippets = " ".join(chunk["text"] for chunk in chunks[:2])
        text = f"Tóm tắt từ corpus demo: {snippets}"

    return Answer(text=text, citations=citations, confidence="Có căn cứ trong corpus")


def synthetic_prefill(profile_id: str) -> dict[str, Any]:
    profiles = {profile["id"]: profile for profile in load_json("sample_profiles.json")}
    return dict(profiles[profile_id])


def parse_uploaded_text(raw: bytes) -> dict[str, Any]:
    text = raw.decode("utf-8", errors="ignore")
    profile: dict[str, Any] = {}
    patterns = {
        "name": r"(?:ten doanh nghiep|tên doanh nghiệp|company)[:\s]+(.+)",
        "tax_code": r"(?:ma so thue|mã số thuế|tax code)[:\s]+([0-9]{8,14})",
        "province": r"(?:tinh|tỉnh|province)[:\s]+(.+)",
        "industry": r"(?:linh vuc|lĩnh vực|industry)[:\s]+(.+)",
        "employees": r"(?:lao dong|lao động|employees)[:\s]+([0-9]+)",
        "revenue_bil": r"(?:doanh thu|revenue)[:\s]+([0-9]+(?:[.,][0-9]+)?)",
        "capital_bil": r"(?:von|vốn|capital)[:\s]+([0-9]+(?:[.,][0-9]+)?)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            value: Any = match.group(1).strip()
            if key in {"employees"}:
                value = int(float(value.replace(",", ".")))
            if key in {"revenue_bil", "capital_bil"}:
                value = float(value.replace(",", "."))
            profile[key] = value

    province_map = {
        "ha noi": "Hà Nội",
        "tp ho chi minh": "TP. Hồ Chí Minh",
        "da nang": "Đà Nẵng",
        "binh duong": "Bình Dương",
        "bac ninh": "Bắc Ninh",
    }
    industry_map = {
        "phan mem / ai": "Phần mềm / AI",
        "phan mem": "Phần mềm / AI",
        "ai": "Phần mềm / AI",
        "san xuat": "Sản xuất",
        "cong nghe cao": "Công nghệ cao",
        "dich vu doi moi sang tao": "Dịch vụ đổi mới sáng tạo",
        "thuong mai": "Thương mại",
    }
    if "province" in profile:
        profile["province"] = province_map.get(normalize_text(str(profile["province"])), profile["province"])
    if "industry" in profile:
        profile["industry"] = industry_map.get(normalize_text(str(profile["industry"])), profile["industry"])

    startup_match = re.search(r"(?:startup|doi moi sang tao|đổi mới sáng tạo)[:\s]+(.+)", text, flags=re.IGNORECASE)
    if startup_match:
        value = normalize_text(startup_match.group(1))
        profile["startup_innovation"] = not any(token in value for token in {"khong", "khong co", "no", "false"})
    elif "startup" in normalize_text(text) or "doi moi sang tao" in normalize_text(text):
        profile["startup_innovation"] = True
    return profile


def generate_dean844_docx(profile: dict[str, Any], selected_policy: dict[str, Any] | None = None) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Đơn đăng ký tham gia nhiệm vụ Đề án 844", level=1)
    document.add_paragraph("Bản demo được điền tự động từ hồ sơ doanh nghiệp trong GrantPilot AI.")

    table = document.add_table(rows=0, cols=2)
    fields = [
        ("Tên doanh nghiệp", profile.get("name", "")),
        ("Mã số thuế", profile.get("tax_code", "")),
        ("Địa phương", profile.get("province", "")),
        ("Lĩnh vực", profile.get("industry", "")),
        ("Ngành nghề/mô tả", profile.get("business_line", "")),
        ("Người đại diện", profile.get("representative", "")),
        ("Email", profile.get("email", "")),
        ("Điện thoại", profile.get("phone", "")),
        ("Số lao động", str(profile.get("employees", ""))),
        ("Doanh thu năm gần nhất", f"{profile.get('revenue_bil', '')} tỷ đồng"),
        ("Vốn", f"{profile.get('capital_bil', '')} tỷ đồng"),
        ("Giai đoạn", profile.get("stage", "")),
    ]
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    document.add_heading("Nội dung đề xuất hỗ trợ", level=2)
    document.add_paragraph(
        "Doanh nghiệp đề xuất được hỗ trợ cố vấn, kết nối hệ sinh thái, hoàn thiện sản phẩm, "
        "rà soát sở hữu trí tuệ và chuẩn hóa hồ sơ tham gia chương trình khởi nghiệp đổi mới sáng tạo."
    )

    document.add_heading("Checklist hồ sơ", level=2)
    checklist = selected_policy.get("checklist", []) if selected_policy else []
    for item in checklist or ["Thuyết minh nhiệm vụ", "Dự toán kinh phí", "Hồ sơ năng lực đơn vị"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Căn cứ demo", level=2)
    for citation in (selected_policy or {}).get("citations", []):
        document.add_paragraph(
            f"{citation['document']} - {citation['clause']} - {citation['status']} - {citation['source']}",
            style="List Bullet",
        )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
